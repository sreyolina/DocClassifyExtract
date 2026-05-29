from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import os

# ─── Generate Flow Diagram Image ─────────────────────────────────────────────
def generate_flow_diagram():
    """Creates a formal, monochrome corporate-style flow diagram for client presentation."""
    fig, ax = plt.subplots(1, 1, figsize=(11, 15))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 15)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # Formal color scheme — lighter tones for readability
    NAVY = '#4A6FA5'
    SLATE = '#5B7FAD'
    CHARCOAL = '#4A4A4A'
    LIGHT_GRAY = '#F0F4F8'
    BORDER = '#3D5A80'
    DECISION_BG = '#FFFFFF'
    STOP_BG = '#FFFFFF'

    def draw_box(x, y, w, h, title, subtitle='', fontsize=11, style='round', filled=True):
        if style == 'diamond':
            diamond = plt.Polygon([[x+w/2, y+h], [x+w, y+h/2], [x+w/2, y], [x, y+h/2]],
                                  closed=True, facecolor=DECISION_BG, edgecolor=BORDER, linewidth=1.5)
            ax.add_patch(diamond)
            ax.text(x + w/2, y + h/2 + 0.1, title, ha='center', va='center',
                    fontsize=fontsize, color=NAVY, fontweight='bold', linespacing=1.3)
            if subtitle:
                ax.text(x + w/2, y + h/2 - 0.22, subtitle, ha='center', va='center',
                        fontsize=fontsize, color=CHARCOAL, fontweight='bold', linespacing=1.2)
        else:
            if filled:
                box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04",
                                     facecolor=NAVY, edgecolor=BORDER, linewidth=1.3)
                ax.add_patch(box)
                text_color = 'white'
            else:
                box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04",
                                     facecolor=LIGHT_GRAY, edgecolor=BORDER, linewidth=1.3)
                ax.add_patch(box)
                text_color = NAVY
            if subtitle:
                ax.text(x + w/2, y + h/2 + 0.15, title, ha='center', va='center',
                        fontsize=fontsize, color=text_color, fontweight='bold', linespacing=1.3)
                ax.text(x + w/2, y + h/2 - 0.18, subtitle, ha='center', va='center',
                        fontsize=fontsize, color=text_color if filled else CHARCOAL,
                        fontweight='bold', linespacing=1.2)
            else:
                ax.text(x + w/2, y + h/2, title, ha='center', va='center',
                        fontsize=fontsize, color=text_color, fontweight='bold',
                        linespacing=1.3)

    def draw_arrow(x1, y1, x2, y2, label='', label_side='right'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=CHARCOAL, lw=1.5,
                                    connectionstyle='arc3,rad=0'))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            offset = 0.15 if label_side == 'right' else -0.15
            ax.text(mx + offset, my + 0.05, label, fontsize=9, color=CHARCOAL,
                    fontweight='bold', ha='center')

    # ── Title ──
    ax.text(5.5, 14.5, 'Document Classification & Extraction Pipeline', ha='center',
            fontsize=17, fontweight='bold', color=NAVY, fontfamily='sans-serif')
    ax.text(5.5, 14.05, 'End-to-End Processing Flow', ha='center',
            fontsize=12, color=SLATE, fontfamily='sans-serif')

    # Thin separator line
    ax.plot([2, 9], [13.9, 13.9], color=SLATE, linewidth=0.5, alpha=0.5)

    # ── Step 1: Document Ingestion ──
    draw_box(2.5, 13.0, 6, 0.7, 'Document Ingestion',
             'PDF upload to Azure Blob Storage\ngenpact/incoming-documents/{name}')
    draw_arrow(5.5, 13.0, 5.5, 12.6)

    # ── Step 2: Event Trigger ──
    draw_box(2.5, 11.8, 6, 0.7, 'Event Trigger',
             'Azure Function  •  ClassifyAndExtract (Blob Trigger)')
    draw_arrow(5.5, 11.8, 5.5, 11.4)

    # ── Step 3: Classification ──
    draw_box(2.5, 10.5, 6, 0.8, 'Document Classification',
             'Azure AI Content Understanding\ndoc_classifier_cre_cni_valuation_confidence_score_other')
    draw_arrow(5.5, 10.5, 5.5, 10.1)

    # ── Step 4: LLM Validation ──
    draw_box(2.5, 9.2, 6, 0.8, 'Classification Validation',
             'GPT-4.1 LLM Verification\n20% page sampling  •  Indicator matching')
    draw_arrow(5.5, 9.2, 5.5, 8.7)

    # ── Step 5: Decision Gate ──
    draw_box(3.5, 7.6, 4, 1.0, 'Confidence \u2265 70%', 'Category \u2260 Other', fontsize=11, style='diamond')

    # NO branch (right)
    draw_arrow(7.5, 8.1, 9.2, 8.1, label='NO', label_side='right')
    draw_box(8.8, 7.8, 2.0, 0.6, 'Human Review', 'Manual classification required',
             fontsize=9.5, filled=False)

    # YES branch (down)
    draw_arrow(5.5, 7.6, 5.5, 7.15)
    ax.text(5.8, 7.35, 'YES', fontsize=10, color=NAVY, fontweight='bold')

    # ── Step 6: Field Extraction ──
    draw_box(2.0, 6.3, 7, 0.75, 'Field Extraction',
             'Schema-driven analysis per document type\nCRE \u2192 cre_loan  |  Valuation \u2192 appraisal_report  |  CNI \u2192 cni_agreement',
             fontsize=10)
    draw_arrow(5.5, 6.3, 5.5, 5.9)

    # ── Step 7: Processing & Enrichment ──
    draw_box(2.0, 5.1, 7, 0.7, 'Field Processing & Enrichment',
             'Extract/Generate methods  •  Citation linking  •  Confidence scoring')
    draw_arrow(5.5, 5.1, 5.5, 4.7)

    # ── Step 8: Data Persistence ──
    draw_box(2.5, 3.9, 6, 0.7, 'Data Persistence',
             'Azure SQL Server\ndbo.FeatureData (MERGE)  •  dbo.JobDetails')
    draw_arrow(5.5, 3.9, 5.5, 3.5)

    # Branch lines
    draw_arrow(5.5, 3.5, 3.5, 3.0)
    draw_arrow(5.5, 3.5, 7.5, 3.0)

    # ── Step 9: SME Assignment ──
    draw_box(1.5, 2.15, 3.5, 0.75, 'SME Assignment',
             'Round-robin  •  Capacity check\ndbo.DocumentAssignment', fontsize=9.5)
    ax.text(3.25, 3.1, 'If review required', fontsize=8, color=CHARCOAL, ha='center',
            style='italic')

    # ── Step 10: Blob Routing ──
    draw_box(6.0, 2.15, 3.5, 0.75, 'Blob Routing',
             'cre_loan/  •  cre_valuation/  •  cni/\nRename \u2192 Delete original', fontsize=9.5)
    ax.text(7.75, 3.1, 'On successful processing', fontsize=8, color=CHARCOAL, ha='center',
            style='italic')

    # ── Completion ──
    draw_arrow(3.25, 2.15, 5.0, 1.5)
    draw_arrow(7.75, 2.15, 6.0, 1.5)
    draw_box(3.75, 0.9, 3.5, 0.55, 'Processing Complete', fontsize=11)

    # ── Minimal legend (no colored boxes — just labeled shapes) ──
    ax.plot([1.5, 9.5], [0.55, 0.55], color=SLATE, linewidth=0.3, alpha=0.4)
    legend_text = 'Rectangle = Process Step     Diamond = Decision Gate     Arrow = Data Flow'
    ax.text(5.5, 0.25, legend_text, ha='center', fontsize=9, color=CHARCOAL,
            fontfamily='sans-serif')

    plt.tight_layout()
    img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data_flow_diagram.png')
    plt.savefig(img_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return img_path

flow_diagram_path = generate_flow_diagram()

doc = Document()

# ─── Styles ───────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── Title Page ───────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Document Classification & Extraction Solution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Intelligent Document Processing Pipeline')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Genpact | Azure AI Solution').font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: May 2026\n').font.size = Pt(12)
info.add_run('Platform: Azure Functions v4 (.NET 8)\n').font.size = Pt(12)
info.add_run('AI Services: Azure Content Understanding + GPT-4.1').font.size = Pt(12)

doc.add_page_break()

# ─── Table of Contents ────────────────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. End-to-End Data Flow Diagram',
    '2. Solution Overview',
    '3. Architecture & Services',
    '4. Document Classification Pipeline',
    '5. GPT-4.1 Confidence Scoring',
    '6. Field Extraction',
    '7. Database Persistence',
    '8. SME Assignment (Human-in-the-Loop)',
    '9. Blob Routing & File Organization',
    '10. Document Categories & Indicators',
    '11. Technology Stack',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ─── 1. End-to-End Data Flow Diagram ──────────────────────────────────────────
doc.add_heading('1. End-to-End Data Flow Diagram', level=1)

# Insert the generated flow diagram image
doc.add_picture(flow_diagram_path, width=Inches(7.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Crisp summary for business audience
doc.add_heading('How It Works', level=2)

summary_points = [
    ('1. Upload', 'Drop a PDF into Azure Blob Storage — processing starts automatically.'),
    ('2. Classify', 'Azure AI identifies the document type: CRE Loan, Valuation Report, or C&I Agreement.'),
    ('3. Validate', 'GPT-4.1 confirms the classification with a confidence score. Documents below 70% are flagged for human review.'),
    ('4. Extract', 'AI extracts key fields (dates, names, amounts, terms) with page-level citations showing exactly where each value was found.'),
    ('5. Store', 'All extracted data is saved to Azure SQL with full audit trail.'),
    ('6. Assign', 'Documents needing review are automatically assigned to the right SME via round-robin.'),
    ('7. Route', 'The PDF is moved to an organized folder and renamed for easy identification.'),
]

for title, desc in summary_points:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ─── 1. Solution Overview ────────────────────────────────────────────────────
doc.add_heading('2. Solution Overview', level=1)

doc.add_paragraph(
    'This solution provides an end-to-end intelligent document processing pipeline that automatically '
    'classifies incoming documents, validates the classification using AI, extracts structured data, '
    'and routes documents to the appropriate teams for review.'
)

doc.add_paragraph()
doc.add_heading('Key Capabilities', level=2)
capabilities = [
    'Automatic document classification into CRE (Commercial Real Estate), CNI (Commercial & Industrial), Valuation, or Other',
    'AI-powered confidence scoring using GPT-4.1 to validate classification accuracy',
    'Structured field extraction using document-type-specific analyzers',
    'Citation mapping with page numbers, bounding boxes, and highlight text',
    'Human-in-the-Loop (HITL) flagging for low-confidence or ambiguous documents',
    'Round-robin SME assignment for documents requiring review',
    'Automated blob routing to classified folders',
    'Full audit trail with job status tracking in SQL Server',
]
for cap in capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_page_break()

# ─── 3. Architecture & Services ────────────────────────────────────────
doc.add_heading('3. Architecture & Services', level=1)

doc.add_paragraph(
    'The solution is built as an Azure Function triggered by blob uploads. It orchestrates '
    'multiple services in a defined sequence:'
)

# Services table
table = doc.add_table(rows=8, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Service', 'Technology', 'Responsibility']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

services_data = [
    ['ContentUnderstandingService', 'Azure AI Content Understanding API', 'Document classification + field extraction via analyzers'],
    ['GPT-4.1 Confidence Scorer', 'Azure OpenAI (GPT-4.1)', 'Validates classification with reasoning and indicator matching'],
    ['DocumentFieldExtractor', 'Custom .NET Service', 'Parses raw API responses into structured fields with citations'],
    ['DatabaseService', 'Azure SQL Server', 'Persists extracted fields, job details, and audit trail'],
    ['SmeAssignmentService', 'Azure SQL Server', 'Assigns documents to SMEs via round-robin for HITL review'],
    ['BlobRoutingService', 'Azure Blob Storage', 'Moves documents to classified folders after processing'],
    ['CitationService', 'Custom .NET Service', 'Maps field values to specific page locations and bounding boxes'],
]
for row_idx, row_data in enumerate(services_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_data

doc.add_paragraph()
doc.add_heading('Trigger Mechanism', level=2)
doc.add_paragraph(
    'The pipeline is triggered automatically when a document is uploaded to the Azure Blob Storage container:'
)
p = doc.add_paragraph()
p.add_run('Container: ').bold = True
p.add_run('genpact')
p = doc.add_paragraph()
p.add_run('Path: ').bold = True
p.add_run('incoming-documents/{filename}')
p = doc.add_paragraph()
p.add_run('Supported Formats: ').bold = True
p.add_run('PDF, DOCX, and other document formats')

doc.add_page_break()

# ─── 4. Document Classification Pipeline ───────────────────────────────
doc.add_heading('4. Document Classification Pipeline', level=1)

doc.add_heading('Step 1: Classifier Provisioning', level=2)
doc.add_paragraph(
    'The system automatically provisions an Azure Content Understanding classifier with the ID:'
)
p = doc.add_paragraph()
p.add_run('doc_classifier_cre_cni_valuation_confidence_score_other').bold = True

doc.add_paragraph(
    'This classifier is configured with detailed category descriptions and strong indicators '
    'for each document type. It uses a pre-built document model as the base analyzer and GPT-4.1 '
    'as the completion model for intelligent classification.'
)

doc.add_heading('Step 2: Document Submission', level=2)
doc.add_paragraph('The document is submitted to the Azure Content Understanding API:')
p = doc.add_paragraph()
p.add_run('API Endpoint: ').bold = True
p.add_run('{endpoint}/contentunderstanding/analyzers/{classifierId}:analyze')
p = doc.add_paragraph()
p.add_run('API Version: ').bold = True
p.add_run('2025-11-01')
p = doc.add_paragraph()
p.add_run('Authentication: ').bold = True
p.add_run('Ocp-Apim-Subscription-Key header')

doc.add_paragraph(
    'The API processes the document asynchronously. The system polls the Operation-Location '
    'URL until classification completes (with exponential backoff, max 900 seconds).'
)

doc.add_heading('Step 3: Classification Results', level=2)
doc.add_paragraph('The classifier returns segments with:')
results = [
    'Category — The document type (CRE, CNI, Valuation, or Other)',
    'Confidence — Native classifier confidence score (0.0 to 1.0)',
    'Start/End Page Numbers — Page range for each classified segment',
    'Markdown Content — Full text extraction for downstream processing',
]
for r in results:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ─── 5. GPT-4.1 Confidence Scoring ────────────────────────────────────
doc.add_heading('5. GPT-4.1 Confidence Scoring', level=1)

doc.add_paragraph(
    'After the initial classification, the system performs a second-level validation using GPT-4.1. '
    'This acts as a quality gate to ensure classification accuracy before proceeding to extraction.'
)

doc.add_heading('How It Works', level=2)
steps = [
    'Extract a representative page sample (20% of document pages) as markdown text',
    'Send the excerpt + assigned category to GPT-4.1 with a structured scoring prompt',
    'GPT-4.1 evaluates whether the document matches the assigned category based on domain-specific indicators',
    'Returns a confidence percentage (0-99) with detailed reasoning',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Scoring Rules', level=2)
scoring_table = doc.add_table(rows=6, cols=2)
scoring_table.style = 'Medium Shading 1 Accent 1'
scoring_table.rows[0].cells[0].text = 'Score Range'
scoring_table.rows[0].cells[1].text = 'Meaning'
scoring_data = [
    ['91-99%', 'Document overwhelmingly matches — many strong indicators present, no competing signals'],
    ['70-89%', 'Document clearly matches — several strong indicators but not all present'],
    ['50-69%', 'Ambiguous — some indicators match but competing signals exist'],
    ['30-49%', 'Weak match — another category may be more appropriate'],
    ['0-29%', 'Likely misclassified'],
]
for row_idx, (score, meaning) in enumerate(scoring_data, start=1):
    scoring_table.rows[row_idx].cells[0].text = score
    scoring_table.rows[row_idx].cells[1].text = meaning

doc.add_paragraph()
doc.add_heading('Confidence Threshold', level=2)
p = doc.add_paragraph()
p.add_run('Threshold: 70%').bold = True
doc.add_paragraph(
    'If the GPT confidence score is below 70%, the document is flagged for human intervention '
    'and extraction is skipped. This ensures that only confidently classified documents proceed '
    'through the automated extraction pipeline.'
)

doc.add_heading('GPT Response Format', level=2)
doc.add_paragraph('GPT-4.1 returns structured JSON with:')
gpt_fields = [
    'confidence_percent — Integer score (0-99)',
    'reasoning — 2-3 sentence explanation of the score',
    'matched_indicators — List of category indicators found in the document',
    'missing_indicators — List of expected indicators not found',
    'competing_category — If another category might be a better fit',
    'competing_confidence_percent — Confidence for the competing category',
]
for f in gpt_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ─── 5. Field Extraction ─────────────────────────────────────────────────────
doc.add_heading('6. Field Extraction', level=1)

doc.add_paragraph(
    'After classification passes the confidence threshold, the system extracts structured fields '
    'using document-type-specific analyzers.'
)

doc.add_heading('Extraction Analyzers', level=2)
analyzer_table = doc.add_table(rows=4, cols=3)
analyzer_table.style = 'Medium Shading 1 Accent 1'
analyzer_table.rows[0].cells[0].text = 'Document Type'
analyzer_table.rows[0].cells[1].text = 'Analyzer ID'
analyzer_table.rows[0].cells[2].text = 'Schema File'
analyzer_data = [
    ['Valuation (Appraisal)', 'appraisal_report_analyzer', 'appraisal_report_analyzer.json'],
    ['CRE (Loan)', 'cre_loan_analyzer', 'cre_loan_analyzer.json'],
    ['CNI (Commercial & Industrial)', 'cni_agreement_analyzer', 'cni_agreement_analyzer.json'],
]
for row_idx, row_data in enumerate(analyzer_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        analyzer_table.rows[row_idx].cells[col_idx].text = cell_data

doc.add_paragraph()
doc.add_heading('Field Methods', level=2)
doc.add_paragraph('Each field in an analyzer schema has a method:')
doc.add_paragraph('Extract — Value is pulled directly from the document text (OCR/layout-based)', style='List Bullet')
doc.add_paragraph('Generate — Value is inferred/summarized by the AI model from context', style='List Bullet')

doc.add_heading('Field-Level Confidence', level=2)
doc.add_paragraph(
    'Each extracted field includes a confidence score. Fields with confidence below 0.5 are flagged '
    'as requiring human review (HITL). Their value is set to "H-I-T-L" to indicate that a subject '
    'matter expert needs to verify/correct the value.'
)

doc.add_heading('Citations', level=2)
doc.add_paragraph('Each extracted field is enriched with citation data:')
citation_fields = [
    'Page number where the value was found',
    'PDF link with SAS token for direct access',
    'Span offset and length in the document text',
    'Highlight text showing the exact source passage',
    'Bounding box coordinates (X, Y, Width, Height) for visual highlighting',
]
for c in citation_fields:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ─── 6. Database Persistence ─────────────────────────────────────────────────
doc.add_heading('7. Database Persistence', level=1)

doc.add_paragraph(
    'All extracted data is persisted to Azure SQL Server using MERGE (upsert) operations.'
)

doc.add_heading('Tables', level=2)

doc.add_heading('FeatureData Table', level=3)
doc.add_paragraph('Stores each extracted field with full metadata:')
feature_fields = [
    'DocumentId — Unique document identifier (derived from filename)',
    'FeatureId — Numeric ID mapped from field name',
    'FieldName — Human-readable field name',
    'Value — Extracted or generated value',
    'Confidence — Field-level confidence (0.0 to 1.0)',
    'FieldMethod — "extract" or "generate"',
    'ConfidenceReason — Explanation of the confidence score',
    'ReviewRequired — Boolean flag for HITL review',
    'Citation fields — PageNumber, PdfLink, SpanOffset, SpanLength, HighlightText, BoundingBox',
]
for f in feature_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('JobDetails Table', level=3)
doc.add_paragraph('Tracks the processing status of each document:')
job_fields = [
    'JobId — Unique processing job identifier',
    'DocumentId — Links to the source document',
    'Status — "Successful", "Partially Successful", "Failed", or "No Fields Extracted"',
]
for f in job_fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ─── 7. SME Assignment ───────────────────────────────────────────────────────
doc.add_heading('8. SME Assignment (Human-in-the-Loop)', level=1)

doc.add_paragraph(
    'When extracted fields require human review (confidence < 0.5), the document is automatically '
    'assigned to a Subject Matter Expert (SME) for verification.'
)

doc.add_heading('Round-Robin Assignment', level=2)
doc.add_paragraph('The assignment algorithm:')
assignment_steps = [
    'Identifies active SMEs for the document type (CRE, CNI, Valuation)',
    'Retrieves the last assigned SME pointer from the RoundRobinPointer table',
    'Selects the next eligible SME (must be active and under concurrent limit)',
    'Creates an assignment record and updates the pointer',
]
for i, step in enumerate(assignment_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Reassignment', level=2)
doc.add_paragraph(
    'Documents can be reassigned to different SMEs if needed. The system marks the current '
    'assignment as "Reassigned" and creates a new assignment for the target SME.'
)

doc.add_heading('When Human Intervention is Required', level=2)
intervention_cases = [
    'Document classified as "Other" (unknown category)',
    'GPT confidence score below 70% for the assigned category',
    'Individual field confidence below 0.5 (field-level HITL)',
]
for case in intervention_cases:
    doc.add_paragraph(case, style='List Bullet')

doc.add_page_break()

# ─── 8. Blob Routing ─────────────────────────────────────────────────────────
doc.add_heading('9. Blob Routing & File Organization', level=1)

doc.add_paragraph(
    'After successful processing, documents are automatically moved from the intake folder '
    'to their classified destination folder.'
)

doc.add_heading('Routing Rules', level=2)
routing_table = doc.add_table(rows=4, cols=2)
routing_table.style = 'Medium Shading 1 Accent 1'
routing_table.rows[0].cells[0].text = 'Document Type'
routing_table.rows[0].cells[1].text = 'Destination Folder'
routing_data = [
    ['CRE (Loan)', 'genpact/cre_loan/'],
    ['Valuation (Appraisal)', 'genpact/cre_valuation/'],
    ['CNI', 'genpact/cni/'],
]
for row_idx, (doc_type, dest) in enumerate(routing_data, start=1):
    routing_table.rows[row_idx].cells[0].text = doc_type
    routing_table.rows[row_idx].cells[1].text = dest

doc.add_paragraph()
doc.add_heading('File Naming Convention', level=2)
doc.add_paragraph(
    'Routed documents are renamed based on extracted metadata (borrower name, document type, date) '
    'to enable easy identification and search.'
)

doc.add_heading('Routing is Skipped When', level=2)
skip_cases = [
    'Document category is "Other"',
    'Confidence score is below 70% (requires human intervention)',
    'Processing encountered errors',
]
for case in skip_cases:
    doc.add_paragraph(case, style='List Bullet')

doc.add_page_break()

# ─── 10. Document Categories & Indicators ───────────────────────────────
doc.add_heading('10. Document Categories & Indicators', level=1)

doc.add_heading('Valuation (Appraisal Reports)', level=2)
doc.add_paragraph(
    'Appraisal or valuation reports focused on estimating property value and collateral value '
    'using valuation methodology and appraisal standards.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Strong Indicators:').bold = True
valuation_indicators = [
    'Effective age, remaining economic life',
    'USPAP compliance, RICS Red Book valuation',
    'Comparable sales analysis, paired sales analysis',
    'Yield capitalization, band of investment',
    'Scope of work, neighborhood analysis',
    'Site coverage ratio, floor area ratio',
    'Going concern value, business enterprise value',
    'Retrospective/prospective valuation',
]
for ind in valuation_indicators:
    doc.add_paragraph(ind, style='List Bullet')

doc.add_heading('CRE (Commercial Real Estate Loans)', level=2)
doc.add_paragraph(
    'Commercial real estate loan agreements or legal credit packages secured by real property, '
    'with covenant, collateral control, tenant controls, and cash-flow control language.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Strong Indicators:').bold = True
cre_indicators = [
    'Assignment of Leases and Rents',
    'Tenant estoppel certificates',
    'Non-recourse carveouts / bad boy guaranty',
    'DSCR triggers, stabilization conditions',
    'Property cash flow waterfall, net cash flow sweep',
    'Mortgage deed / deed of trust',
    'Environmental indemnity agreement',
    'Occupancy covenants, minimum occupancy requirement',
]
for ind in cre_indicators:
    doc.add_paragraph(ind, style='List Bullet')

doc.add_heading('CNI (Commercial & Industrial)', level=2)
doc.add_paragraph(
    'Commercial and industrial loan agreements, corporate credit agreements, borrowing base documents, '
    'and general business lending packages not primarily CRE or valuation.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Strong Indicators:').bold = True
cni_indicators = [
    'Revolving and term loan mechanics',
    'Corporate financial covenants',
    'Borrower and guarantor corporate credit terms',
    'Collateral structures not centered on real estate',
    'Borrowing base calculations',
]
for ind in cni_indicators:
    doc.add_paragraph(ind, style='List Bullet')

doc.add_heading('Other', level=2)
doc.add_paragraph(
    'Documents that do not clearly fall into CRE, CNI, or Valuation. These are automatically '
    'flagged for human intervention with no extraction performed.'
)

doc.add_page_break()

# ─── 11. Technology Stack ─────────────────────────────────────────────────────
doc.add_heading('11. Technology Stack', level=1)

tech_table = doc.add_table(rows=9, cols=2)
tech_table.style = 'Medium Shading 1 Accent 1'
tech_table.rows[0].cells[0].text = 'Component'
tech_table.rows[0].cells[1].text = 'Technology'
tech_data = [
    ['Runtime', 'Azure Functions v4 (.NET 8 Isolated Worker)'],
    ['Classification', 'Azure AI Content Understanding (API version 2025-11-01)'],
    ['Confidence Scoring', 'Azure OpenAI GPT-4.1 (API version 2024-12-01-preview)'],
    ['Storage', 'Azure Blob Storage'],
    ['Database', 'Azure SQL Server'],
    ['Authentication', 'API Key (Ocp-Apim-Subscription-Key)'],
    ['Trigger', 'Blob Trigger (automatic on upload)'],
    ['Language', 'C# 12 / .NET 8'],
]
for row_idx, (component, tech) in enumerate(tech_data, start=1):
    tech_table.rows[row_idx].cells[0].text = component
    tech_table.rows[row_idx].cells[1].text = tech

doc.add_paragraph()
doc.add_heading('Azure Resources', level=2)
resources = [
    'Azure AI Services (Content Understanding + OpenAI) — doc-classify-foundry.cognitiveservices.azure.com',
    'Azure Storage Account — storagegenpactmvp2',
    'Azure SQL Database — genpactpoc on docstream.database.windows.net',
    'Azure Functions App — DocClassifyExtract',
]
for r in resources:
    doc.add_paragraph(r, style='List Bullet')

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\somolinasaha\Downloads\Genpact\DocClassifyExtract\DocClassifyExtract\DocClassifyExtract_Solution_Overview.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
