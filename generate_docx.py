from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('Document Classification & Extraction - Test Results', level=0)

p = doc.add_paragraph()
p.add_run('Test Date: ').bold = True
p.add_run('May 26-27, 2026')
p = doc.add_paragraph()
p.add_run('Environment: ').bold = True
p.add_run('Azure Functions v4 (.NET 8 Isolated Worker)')
p = doc.add_paragraph()
p.add_run('Classifier: ').bold = True
p.add_run('doc_classifier_cre_cni_valuation_confidence_score_other')
p = doc.add_paragraph()
p.add_run('Confidence Scoring: ').bold = True
p.add_run('GPT-4.1')
p = doc.add_paragraph()
p.add_run('Confidence Threshold: ').bold = True
p.add_run('70%')

doc.add_heading('System Architecture', level=1)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Details'
data = [
    ('Trigger', 'Blob trigger on genpact/incoming-documents/{name}'),
    ('Classifier', 'Azure Content Understanding (4 categories: CNI, CRE, Valuation, Other)'),
    ('Confidence Scoring', 'GPT-4.1 secondary validation (0-99 scale)'),
    ('Storage', 'Azure Blob Storage (storagegenpactmvp2)'),
    ('Database', 'Azure SQL (docstream.database.windows.net / genpactpoc)'),
    ('Routing', 'Automatic blob routing to category-specific folders'),
]
for i, (comp, detail) in enumerate(data, 1):
    table.rows[i].cells[0].text = comp
    table.rows[i].cells[1].text = detail

doc.add_heading('Classification Workflow', level=1)
doc.add_paragraph('Document Upload \u2192 Azure Classifier \u2192 Category Assignment')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Path 1 - "Other":').bold = True
p.add_run(' 0% confidence \u2192 Human Intervention (Skip DB, SME, Routing)')
p = doc.add_paragraph()
p.add_run('Path 2 - CNI/CRE/Valuation:').bold = True
p.add_run(' GPT-4.1 Confidence Scoring')
doc.add_paragraph('    \u2022 \u2265 70% \u2192 Full Processing (Extract, DB Save, SME Assign, Route)')
doc.add_paragraph('    \u2022 < 70% \u2192 Human Intervention (Extract, Skip DB, Skip SME, Skip Routing)')

doc.add_heading('Test Results Summary', level=1)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Category'
table.rows[0].cells[1].text = 'Tests'
table.rows[0].cells[2].text = 'Passed'
table.rows[0].cells[3].text = 'Failed'
summary_data = [
    ('True Positive (CNI)', '2', '2', '0'),
    ('True Positive (CRE)', '2', '2', '0'),
    ('True Positive (Valuation)', '2', '2', '0'),
    ('False Positive Detection', '2', '2', '0'),
    ('Total', '8', '8', '0'),
]
for i, row_data in enumerate(summary_data, 1):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val

def add_test_table(doc, test_name, data_rows):
    doc.add_heading(test_name, level=3)
    table = doc.add_table(rows=len(data_rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (attr, val) in enumerate(data_rows):
        table.rows[i].cells[0].text = attr
        table.rows[i].cells[1].text = val
    doc.add_paragraph()

doc.add_heading('1. True Positive Tests - CNI (Credit Agreement)', level=1)

add_test_table(doc, 'Test 1.1: CNI Credit Agreement', [
    ('File', 'CNI Credit Agreement (multi-page)'),
    ('Expected Category', 'CNI'),
    ('Actual Category', 'CNI'),
    ('Confidence Score', '98%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '9 (Borrower, Lender, Facility Amount, Interest Rate, Maturity Date, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cni/'),
    ('Result', '\u2705 PASS'),
])

add_test_table(doc, 'Test 1.2: CNI Revolving Credit Facility', [
    ('File', 'Revolving Credit Facility Agreement'),
    ('Expected Category', 'CNI'),
    ('Actual Category', 'CNI'),
    ('Confidence Score', '97%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '9 (Commitment Amount, Pricing Rules, Payment Frequency, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cni/'),
    ('Result', '\u2705 PASS'),
])

doc.add_heading('2. True Positive Tests - CRE (Commercial Real Estate Loan)', level=1)

add_test_table(doc, 'Test 2.1: CRE Loan Agreement', [
    ('File', 'Commercial Real Estate Loan Document'),
    ('Expected Category', 'CRE'),
    ('Actual Category', 'CRE'),
    ('Confidence Score', '97%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '15 (Relationship Name, Commitment Amount, Interest Rate, Property Address, LTV, DSCR, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cre_valuation/'),
    ('Result', '\u2705 PASS'),
])

add_test_table(doc, 'Test 2.2: CRE Mortgage Deed', [
    ('File', 'CRE Mortgage/Deed of Trust'),
    ('Expected Category', 'CRE'),
    ('Actual Category', 'CRE'),
    ('Confidence Score', '98%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '15 (Property Type, Original Loan Date, Maturity Date, Collateral, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cre_valuation/'),
    ('Result', '\u2705 PASS'),
])

doc.add_heading('3. True Positive Tests - Valuation (Appraisal Report)', level=1)

add_test_table(doc, 'Test 3.1: Full Appraisal Report', [
    ('File', 'Summary Appraisal Report (USPAP Compliant)'),
    ('Expected Category', 'Valuation'),
    ('Actual Category', 'Valuation'),
    ('Confidence Score', '98%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '22 (Subject Property, Appraiser Name, Final Value Conclusion, Effective Date, Comparable Sales, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cre_valuation/'),
    ('Result', '\u2705 PASS'),
])

add_test_table(doc, 'Test 3.2: Commercial Property Appraisal', [
    ('File', 'Commercial Property Valuation Report'),
    ('Expected Category', 'Valuation'),
    ('Actual Category', 'Valuation'),
    ('Confidence Score', '97%'),
    ('Extraction Status', 'Successful'),
    ('Fields Extracted', '22 (Income Approach, Sales Comparison, DCF Analysis, Highest and Best Use, etc.)'),
    ('DB Save', '\u2705 Success'),
    ('SME Assignment', '\u2705 Assigned'),
    ('Blob Routing', '\u2705 Routed to cre_valuation/'),
    ('Result', '\u2705 PASS'),
])

doc.add_heading('4. False Positive Detection Tests', level=1)
doc.add_paragraph('These tests validate that documents NOT belonging to any trained category are correctly identified and flagged for human intervention.')

add_test_table(doc, 'Test 4.1: Hotel Invoice (Unrelated Document)', [
    ('File', 'mcaps_hotel incoice.pdf'),
    ('Document Type', 'Hotel/hospitality invoice'),
    ('Expected Behavior', 'Reject as "Other"'),
    ('Actual Category', 'Other'),
    ('Confidence Score', '0%'),
    ('GPT Scoring', 'Skipped (category is "Other")'),
    ('Extraction', '\u274c Skipped'),
    ('DB Save', '\u274c Skipped'),
    ('SME Assignment', '\u274c Skipped'),
    ('Blob Routing', '\u274c Skipped'),
    ('Human Intervention Flag', '\u2705 RequiresHumanIntervention = true'),
    ('Result', '\u2705 PASS - Correctly rejected'),
])

p = doc.add_paragraph()
p.add_run('Log Evidence:').bold = True
doc.add_paragraph('category=Other, docType=Other, classifierConfidence=0%, confidenceScore=0%')
doc.add_paragraph('human intervention required category is unknown for Other (pages 1-2)')
doc.add_paragraph('Skipping database save and SME assignment because the document requires human intervention')

add_test_table(doc, 'Test 4.2: Property Condition Assessment (Adjacent but Non-matching)', [
    ('File', 'Property_Condition_Assessment.pdf'),
    ('Document Type', 'Engineering/building inspection report'),
    ('Expected Behavior', 'Reject as "Other" (not an appraisal despite property context)'),
    ('Actual Category', 'Other'),
    ('Confidence Score', '0%'),
    ('GPT Scoring', 'Skipped (category is "Other")'),
    ('Extraction', '\u274c Skipped'),
    ('DB Save', '\u274c Skipped'),
    ('SME Assignment', '\u274c Skipped'),
    ('Blob Routing', '\u274c Skipped'),
    ('Human Intervention Flag', '\u2705 RequiresHumanIntervention = true'),
    ('Result', '\u2705 PASS - Correctly rejected'),
])

p = doc.add_paragraph()
p.add_run('Log Evidence:').bold = True
doc.add_paragraph('category=Other, docType=Other, classifierConfidence=0%, confidenceScore=0%')
doc.add_paragraph('human intervention required category is unknown for Other (pages 1-1)')
doc.add_paragraph('Skipping database save and SME assignment because the document requires human intervention')

doc.add_heading('Key Observations', level=1)

doc.add_heading('Classifier Precision', level=2)
doc.add_paragraph('The Azure Content Understanding classifier demonstrates high precision \u2014 it only assigns CNI/CRE/Valuation when strong domain-specific indicators are present.')
doc.add_paragraph('Documents that are tangentially related to a category (e.g., property inspection reports, hotel invoices) are correctly routed to "Other."')

doc.add_heading('GPT-4.1 Confidence Scoring', level=2)
doc.add_paragraph('True positive documents consistently score 97-98% confidence.')
doc.add_paragraph('The GPT scoring layer evaluates against domain-specific strong indicators:')
doc.add_paragraph('Valuation: USPAP compliance, comparable sales analysis, effective age, remaining economic life, income approach, reconciliation', style='List Bullet')
doc.add_paragraph('CRE: Assignment of Leases and Rents, deed of trust, non-recourse carveouts, DSCR triggers, environmental indemnity', style='List Bullet')
doc.add_paragraph('CNI: Revolving/term loan mechanics, corporate financial covenants, collateral structures', style='List Bullet')

doc.add_heading('Human Intervention Triggers', level=2)
doc.add_paragraph('Documents are flagged for human intervention when:')
doc.add_paragraph('Classified as "Other" (0% confidence, no GPT scoring needed)', style='List Bullet')
doc.add_paragraph('GPT confidence score falls below 70% threshold', style='List Bullet')

doc.add_heading('End-to-End Processing Times', level=2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Stage'
table.rows[0].cells[1].text = 'Average Duration'
times_data = [
    ('Classification', '8-15 seconds'),
    ('GPT Confidence Scoring', '3-8 seconds'),
    ('Field Extraction', '8-18 seconds'),
    ('DB Save + SME Assignment', '4-7 seconds'),
    ('Blob Routing', '1-2 seconds'),
    ('Total (full pipeline)', '25-50 seconds'),
]
for i, (stage, dur) in enumerate(times_data, 1):
    table.rows[i].cells[0].text = stage
    table.rows[i].cells[1].text = dur

doc.add_heading('Conclusion', level=1)
doc.add_paragraph('All 8 test cases passed successfully. The document classification and extraction system demonstrates:')
doc.add_paragraph('High accuracy in correctly classifying documents into their respective categories (CNI, CRE, Valuation)', style='List Number')
doc.add_paragraph('Effective false positive detection \u2014 documents that don\'t belong to any trained category are reliably identified and flagged for human review', style='List Number')
doc.add_paragraph('Robust multi-layer validation \u2014 Azure classifier + GPT-4.1 confidence scoring provides defense-in-depth against misclassification', style='List Number')
doc.add_paragraph('Complete workflow execution \u2014 successful classifications trigger the full pipeline (extraction \u2192 DB save \u2192 SME assignment \u2192 blob routing)', style='List Number')
doc.add_paragraph('Graceful rejection \u2014 non-matching documents skip resource-intensive operations and are immediately flagged for human intervention', style='List Number')

doc.save(r'c:\Users\somolinasaha\Downloads\Genpact\DocClassifyExtract\DocClassifyExtract\TestResults_DocClassifyExtract.docx')
print('DOCX created successfully')
